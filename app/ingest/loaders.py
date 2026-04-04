from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from llama_index.core import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}


def load_source_documents(data_directory: Path) -> list[Document]:
    """Load supported local files into LlamaIndex documents."""

    if not data_directory.exists():
        raise FileNotFoundError(
            f"Raw data directory does not exist: `{data_directory}`."
        )

    source_files = sorted(
        path
        for path in data_directory.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )

    documents: list[Document] = []
    for source_file in source_files:
        text = _read_file_text(source_file)
        if not text.strip():
            continue

        relative_path = source_file.relative_to(data_directory)
        documents.append(
            Document(
                text=text,
                metadata={
                    "file_name": source_file.name,
                    "source_path": str(relative_path),
                    "file_extension": source_file.suffix.lower(),
                },
            )
        )

    return documents


def _read_file_text(source_file: Path) -> str:
    extension = source_file.suffix.lower()
    if extension in {".md", ".txt"}:
        return source_file.read_text(encoding="utf-8")
    if extension == ".pdf":
        return _read_pdf_text(source_file)
    if extension == ".docx":
        return _read_docx_text(source_file)

    raise ValueError(f"Unsupported file type: `{extension}`.")


def _read_pdf_text(source_file: Path) -> str:
    reader = PdfReader(source_file)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _read_docx_text(source_file: Path) -> str:
    document = DocxDocument(source_file)
    return "\n".join(paragraph.text for paragraph in document.paragraphs)
